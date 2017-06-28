# -*- coding: utf-8 -*-
import sys
import os
import cx_Oracle
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from optparse import OptionParser
import datetime

def printf (format,*args):
  sys.stdout.write (format % args)

def printException (exception):
  error, = exception.args
  printf ("Error code = %s\n",error.code);
  printf ("Error message = %s\n",error.message);
  
def set_col_widths(table):
    widths = (Inches(0.5), Inches(2.5), Inches(1.5),Inches(0.5),Inches(5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


parser = OptionParser(usage="usage: %prog [options] progtype progname",
                          version="%prog 1.0")
parser.add_option("-c", "--connstr",
                      action="store",
                      dest="connstr",
                      default="apps/appspwd@DB",
                      help="apps/appspwd@DBSID")

parser.add_option("-l", "--list",
                      action="store",
                      dest="listfile",
                      default="C:\list.txt",
                      help="table list.")
 
parser.add_option("-o", "--output",
                      action="store",
                      dest="outputfile",
                      default="C:\schema"+datetime.datetime.now().strftime("%Y%m%d%H%M%S"),
                      help="output files name.")

parser.add_option("-s", "--oraclehome",
                      action="store",
                      dest="oracle",
                      default="C:\ORACLE_HOME",
                      help="oracle home.")


(options, args) = parser.parse_args()
printf ('****************************************************\n')
printf ('Connect String : '+options.connstr+'\n')
printf ('Table List     : '+options.listfile+'\n')
printf ('Output File    : '+options.outputfile+'.docx and .sql\n')
printf ('Oracle Home    : '+options.oracle+'\n')
printf ('****************************************************\n\n')

username = options.connstr.split('/',1)[0]
password = options.connstr.split('/',1)[1].split('@',1)[0]
databaseName = options.connstr.split('/',1)[1].split('@',1)[1]

os.environ["ORACLE_HOME"] = options.oracle

try:
  connection = cx_Oracle.connect (username,password,databaseName)
except cx_Oracle.DatabaseError, exception:
  printf ('Failed to connect to %s\n',databaseName)
  printException (exception)
  exit (1)

cursor = connection.cursor()
cur = connection.cursor()
document = Document()

s = open(options.outputfile+".sql",'w') 

with open(options.listfile,'r') as f:
    for line in f:
        line = line.strip('\n')
        line = line.strip('\t')
        tabName=line.upper()

        try:
          cursor.execute("""
          SELECT TO_CHAR(ATC.COLUMN_ID) || '.',
                 ATC.COLUMN_NAME,
                 ATC.DATA_TYPE,
                 DECODE(ATC.DATA_TYPE, 'VARCHAR2', TO_CHAR(ATC.DATA_LENGTH), ''),
                 ATC.NULLABLE,
                 ACC.COMMENTS
            FROM ALL_TAB_COLUMNS ATC, ALL_COL_COMMENTS ACC
           WHERE ATC.TABLE_NAME = :tableName
             AND ATC.TABLE_NAME = ACC.TABLE_NAME
             AND ATC.COLUMN_NAME = ACC.COLUMN_NAME
           ORDER BY ATC.COLUMN_ID """, 
           {"tableName":tabName})
          printf('======================= '+tabName+' =====================\n')
          table1 = document.add_table(1,1)
          table1.style = 'Table Grid'
          
          t1 = table1.cell(0,0)
          t1.text='Table Name: '+ tabName
          shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          t1._tc.get_or_add_tcPr().append(shading_elm)
          
          table2 = document.add_table(1,5)
          table2.style = 'Table Grid'
          t2 = table2.cell(0,0)
          t3 = table2.cell(0,1)
          t4 = table2.cell(0,2)
          t5 = table2.cell(0,3)
          t6 = table2.cell(0,4)
          
          shading_elm2 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          shading_elm3 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          shading_elm4 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          shading_elm5 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          shading_elm6 = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
          #t2.text=unicode('順序', 'utf-8')
          t2.text=u'順序'
          t3.text=u'欄位名稱'
          t4.text=u'資料型態'
          t5.text=u'欄位\n長度'
          t6.text=u'說明'
          t2.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
          t3.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
          t4.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
          t5.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
          t6.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
          t2._tc.get_or_add_tcPr().append(shading_elm2)
          t3._tc.get_or_add_tcPr().append(shading_elm3)
          t4._tc.get_or_add_tcPr().append(shading_elm4)
          t5._tc.get_or_add_tcPr().append(shading_elm5)
          t6._tc.get_or_add_tcPr().append(shading_elm6)
          set_col_widths(table2)

          s.write('/***  '+tabName+' ***/\n')
          s.write('CREATE TABLE '+tabName + '( \n')
          first = True 
          cstr = ""
          nullstr = ""

          rows = cursor.fetchall()
          for colNum, colName, dataType, dataLen, nullable, cmmtStr in rows:
             if first:
                first=False
             else:
                s.write(', \n')
             table3 = document.add_table(1,5)
             table3.style = 'Table Grid'
             t2 = table3.cell(0,0)
             t3 = table3.cell(0,1)
             t4 = table3.cell(0,2)
             t5 = table3.cell(0,3)
             t6 = table3.cell(0,4)
                       
             t2.text=colNum
             #t2.paragraphs[0].style='ListNumber'
             t3.text=colName
             t4.text=dataType
             if nullable == "Y":
                nullstr = ""
             else:
                nullstr = " NOT NULL"
             if dataLen is None:
                t5.text=''
                s.write(colName + '  '+dataType+nullstr)
             else:
                t5.text=dataLen
                s.write(colName + '  '+dataType+'('+dataLen+')'+nullstr)
             t5.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
             if cmmtStr is None:
                t6.text=''
             else:
                t6.text=cmmtStr
                cstr=cstr+"COMMENT ON COLUMN "+tabName+"."+colName+"\n IS '"+cmmtStr+"';\n"
             set_col_widths(table3)
             #printf (colName + ' '+dataType+'\n')
          p = document.add_paragraph('')
          s.write('\n);\n\n')

        except cx_Oracle.DatabaseError, exception:
          printf ('Failed to execute sql 001\n')
          printException (exception)
          exit (1)
        s.write(cstr+"\n")
    
        try:
          cursor.execute(""" 
          SELECT 'CREATE ' || DECODE(UNIQUENESS, 'UNIQUE ', 'UNIQUE', '') || 'INDEX ' ||
                 INDEX_NAME || ' ON ' || TABLE_NAME || '( ',
                 INDEX_NAME
            FROM ALL_INDEXES
           WHERE TABLE_NAME = :tName
           ORDER BY INDEX_NAME """, 
           {"tName" : str(tabName)})
          rows = cursor.fetchall()
          for idxstr,idxName in rows:
            s.write(idxstr)
            try:
                cur.execute(""" 
                SELECT COLUMN_NAME, TABLE_OWNER
                  FROM ALL_IND_COLUMNS
                 WHERE INDEX_NAME = :indexName
                 ORDER BY COLUMN_POSITION """, 
                 {"indexName":idxName})
                idlfirst = True
                rs = cur.fetchall()
                for iclName,dummy in rs:
                    if idlfirst:
                        s.write(iclName )
                        idlfirst = False
                    else:
                        s.write(","+iclName)
                s.write(" );\n")
            except cx_Oracle.DatabaseError, exception:
                              printf ('Failed to execute sql 003\n')
                              printException (exception)
                              exit (1)          
        except cx_Oracle.DatabaseError, exception:
                  printf ('Failed to execute sql 002\n')
                  printException (exception)
                  exit (1)
        
        s.write('\n\n\n')
        document.add_page_break()

document.save(options.outputfile+".docx")
s.close()

cursor.close()
cur.close()

connection.close()

exit (0)